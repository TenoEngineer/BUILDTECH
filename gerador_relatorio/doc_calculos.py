# -*- coding: utf-8 -*-
"""
Created on Fri Sep 10 12:22:13 2021

@author: heitor

RELATÓRIO DE CÁLCULOS

"""

import re
import sys
from os import path
from pathlib import Path
from tkinter.font import BOLD

from docx import Document
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

import root_var as rt


def calculos():

    validacao = rt.validacao_calculos

    if validacao is True:
        path_calcu = rt.caminho
        city = rt.cidade
        path_path = Path(path_calcu)
        calculos = list(path_path.glob('**\*.png'))
        calculos = sorted(calculos, key=lambda x: [int(
            k) if k.isdigit() else k for k in re.split('([0-9]+)', x.stem)])
        modelo_calculos = 'MODELO_CALCULOS.docx'
        calculos_total = 'CALCULOS_TOTAL.docx'

        if getattr(sys, 'frozen', False):
            application_path = path.dirname(sys.executable)
        elif __file__:
            application_path = path.dirname(__file__)

        doc_calculos = Document(path.join(application_path, modelo_calculos))

        for i, image in enumerate(calculos):
            picture = str(image)
            p = doc_calculos.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(picture, width=Inches(10.8), height=Inches(6.6))
            name_poste = path.split(image)
            name = (name_poste[1].split('.'))[0]
            text = p.add_run(f'POSTE {name}')
            text.bold = True
            text.font.size = Pt(18)
            #p.add_run(f'POSTE {name}').bold = True
            #doc_calculos.add_paragraph(f'POSTE {name}').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc_calculos.save(path.join(application_path, calculos_total))

        convert(path.join(application_path, calculos_total),
                path.join(path_calcu, f'{city}_CALCULOS.pdf'))

        return True
