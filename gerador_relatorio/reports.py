import asyncio
import os
import re
import sys
from glob import glob
from os import path
from pathlib import Path
from tkinter.font import BOLD

from docx import Document
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from numpy import tile
from pdf2image import convert_from_path
from PIL import Image
from PyPDF2 import PdfFileReader, PdfFileWriter

import documentacao as dc
import root_var as rt

if getattr(sys, 'frozen', False):
    application_path = path.dirname(sys.executable)
elif __file__:
    application_path = path.dirname(__file__)


async def calculations():
    # 3 - Só começa a rodar depois que rodou pdftopng
    await pdftopng()
    path_calcu = rt.caminho
    city = rt.cidade
    path_path = Path(path_calcu)
    calculos = list(path_path.glob('**\*.png'))
    calculos = sorted(calculos, key=lambda x: [int(
        k) if k.isdigit() else k for k in re.split('([0-9]+)', x.stem)])
    modelo_calculos = 'MODELO_CALCULOS.docx'
    calculos_total = 'CALCULOS_TOTAL.docx'

    if getattr(sys, 'frozen', False):
        application_path = path.dirname(sys.executable)
    elif __file__:
        application_path = path.dirname(__file__)

    doc_calculos = Document(path.join(application_path, modelo_calculos))

    for image in enumerate(calculos):
        p = doc_calculos.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image[1]), width=Inches(10.8), height=Inches(6.6))
        name_poste = path.split(image[1])
        name = (name_poste[1].split('.'))[0]
        text = p.add_run(f'POSTE {name}')
        text.bold = True
        text.font.size = Pt(18)

    doc_calculos.save(path.join(application_path, calculos_total))

    convert(path.join(application_path, calculos_total),
            path.join(path_calcu, f'{city}_CALCULOS.pdf'))


async def photograph():
    path_fotos = rt.caminho
    city = rt.cidade
    path_path = Path(path_fotos)
    fotos = list(path_path.glob('**\*.jpg'))
    relatorio_fotografico = 'RELATORIO FOTOGRAFICO MODELO A4.docx'
    fotos_total = 'FOTOS_TOTAL.docx'

    fotos = sorted(fotos, key=lambda x: [
        int(k) if k.isdigit() else k for k in re.split('([0-9]+)', x.stem)])
    relatorio = str(
        path.join(application_path, relatorio_fotografico))
    doc_fotos = Document(relatorio)

    await editonDocx(doc_fotos)

    document = Document(path.join(application_path, 'FOTOS.docx'))

    for image in enumerate(fotos):
        picture = path.join(path_fotos, str(image[1]))
        p = document.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(picture, width=Inches(10), height=Inches(6))

    document.save(path.join(application_path, fotos_total))

    city_fotos = f'{city}_FOTOS.pdf'

    convert(path.join(application_path, fotos_total),
            path.join(path_fotos, city_fotos))


async def editonDocx(document):

    cidade = rt.cidade.upper()
    fotos = 'FOTOS.docx'
    for section in document.sections:
        header = section.header
        header.paragraphs[1].text = 'LOCALIDADE: CIDADE - RS'
        header.paragraphs[1].text = header.paragraphs[1].text.replace(
            "CIDADE", cidade)
        header.paragraphs[1].bold = True

    document.save(path.join(application_path, fotos))


async def split():

    await calculations()
    await photograph()

    caminho = rt.caminho
    pdfs = glob(path.join(caminho, "*.pdf"))

    for i in pdfs:

        pdf = PdfFileReader(open(str(i), "rb"))
        size_pdf = path.getsize(str(i))
        target_size_limit = 9
        target_size = target_size_limit*1024*1024
        pages = pdf.getNumPages()

        if size_pdf > target_size:  # SE O ARQUIVO FOR MAIOR DE 9 MB

            list_sizes = []

            for page in range(pages):  # SELECIONA TODAS AS PAGINAS
                output = PdfFileWriter()
                output.addPage(pdf.getPage(int(page)))
                file_name = path.join(
                    caminho, f'{page}.pdf')  # ARQUIVO DE UMA PAGE
                with open(file_name, 'wb') as out:  # CRIA NOVO ARQUIVO COM UMA PAGINA
                    output.write(out)
                size_page = path.getsize(
                    file_name)  # TAMANHO DO ARQUIVO
                list_sizes.append(size_page)  # ADICIONA NA LISTA O TAMANHO

            partition_weight = 0
            list_pages_merge = [[]]

            # PEGA LISTA DO TAMANHO DAS PAGE
            for m, weight in enumerate(list_sizes):
                if partition_weight + weight < target_size:  # SOMA COM A PRÓXIMA PAGE E VERIFICA SE FICA MAIOR QUE 9MB
                    partition_weight += weight  # FAZ A SOMA DELES
                    # CRIA UMA LISTA DO TOTAL DE PAGINAS QUE FICARÃO EM UM UNICOS PDF
                    list_pages_merge[-1].append(m)
                else:
                    partition_weight = weight  # SE SOMAR MAIOR QUE 9MB FAZ COM QUE CONTINUE O LOOP
                    # ADICIONA A LISTA EM UMA LISTA
                    list_pages_merge.append([m])

            # LOOP PARA A GERAÇÃO DE PDFS
            for partition, pages in enumerate(list_pages_merge):
                # NOME DO NOVO PDF: CIDADE_FOTOS_X
                file_name = f'{str(i)[0:-4]}_{partition}.pdf'
                output = PdfFileWriter()
                for page in list_pages_merge[partition]:
                    pdf_marge = PdfFileReader(
                        path.join(caminho, f'{page}.pdf'))
                    output.addPage(pdf_marge.getPage(0))
                with open(file_name, 'wb') as out:  # CRIA NOVO ARQUIVO COM UMA PAGINA
                    output.write(out)

            for page in range(pdf.getNumPages()):
                os.remove(path.join(caminho, f'{page}.pdf'))


async def jpgtopng():
    # 1 - Converto a pasta de calculos os arquivos que se encontram em jpg para png

    # procurar a pasta dos calculos
    caminho = rt.caminho
    calcu = os.listdir(caminho)
    for calculo in calcu:
        if calculo.startswith("CALCULOS"):
            caminho_calculos = os.path.join(caminho, calculo)
    pictures = os.listdir(caminho_calculos)
    pictures = list(filter(lambda picture: '.jpg' in picture, pictures))

    for pic in pictures:
        print(f'Gerado arquivo: {pic}')
        image = Image.open(os.path.join(caminho_calculos, pic))
        name = pic[:-4]
        new_name = os.path.join(caminho_calculos, str(name))
        image.save(f'{new_name}.png')
        os.remove(os.path.join(caminho_calculos, str(pic)))
    print('Conversão JPG - PNG finalizada')


async def pdftopng():
    # 2 - Converto a pasta de calculos os arquivos que se encontram em pdf para png
    # Roda só depois de jpgtopng
    await jpgtopng()
    # procurar pasta com as fotos
    poppler = 'poppler-0.68.0/bin'
    user_path = os.path.expanduser('~')
    poppler_path = os.path.join(user_path, poppler)

    caminho = rt.caminho
    calcu = os.listdir(caminho)
    for calculo in calcu:
        if calculo.startswith("CALCULOS"):
            caminho_calculos = os.path.join(caminho, calculo)
    pictures = os.listdir(caminho_calculos)
    pictures = list(filter(lambda picture: '.pdf' in picture, pictures))

    for pic in pictures:
        image = convert_from_path(os.path.join(
            caminho_calculos, pic), poppler_path=poppler_path)
        images = list(range(len(image)))
        name_pdf = pic.replace('POSTE', '')
        name_pdf = name_pdf.replace(' ', '')
        name_png = name_pdf.replace('.pdf', '.png')
        image = image[0]
        image.save(os.path.join(caminho_calculos, name_png), 'PNG')
        os.remove(os.path.join(caminho_calculos, pic))
        print(f'Gerado arquivo: {name_png}')
    print('Conversão PDF - PNG finalizada')


def main():
    asyncio.run(split())


if __name__ == "__main__":
    main()
