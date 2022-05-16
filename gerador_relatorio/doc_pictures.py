# -*- coding: utf-8 -*-
"""
Created on Wed Sep  8 11:35:01 2021

@author: heitor

RELATÓRIO PDF
"""

import re
import sys
from os import path
from pathlib import Path

from docx import Document
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

import documentacao as dc
import root_var as rt


def fotos():

    validacao = rt.validacao_fotos

    if validacao is True:
        path_fotos = rt.caminho
        city = rt.cidade
        path_path = Path(path_fotos)
        fotos = list(path_path.glob('**\*.jpg'))
        relatorio_fotografico = 'RELATORIO FOTOGRAFICO MODELO A4.docx'
        fotos_total = 'FOTOS_TOTAL.docx'

        fotos = sorted(fotos, key=lambda x: [
            int(k) if k.isdigit() else k for k in re.split('([0-9]+)', x.stem)])

        if getattr(sys, 'frozen', False):
            application_path = path.dirname(sys.executable)
        elif __file__:
            application_path = path.dirname(__file__)

        relatorio = str(
            path.join(application_path, relatorio_fotografico))
        doc_fotos = Document(relatorio)

        dc.documentacao(doc_fotos)

        document = Document(path.join(application_path, 'FOTOS.docx'))

        for i, image in enumerate(fotos):
            picture = path.join(path_fotos, image)
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(picture, width=Inches(10), height=Inches(6))

        document.save(path.join(application_path, fotos_total))

        city_fotos = f'{city}_FOTOS.pdf'
        convert(path.join(application_path, fotos_total),
                path.join(path_fotos, city_fotos))

        return True
