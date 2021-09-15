# -*- coding: utf-8 -*-
"""
Created on Fri Sep 10 12:22:13 2021

@author: heitor

RELATÓRIO DE CÁLCULOS

"""

from docx import Document
import document as dc
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from pathlib import Path
from docx.shared import Inches
import re
import doc_pictures

path_calcu = doc_pictures.path_fotos
city = dc.cidade
path = Path(path_calcu)
calculos = list(path.glob('**\*.png'))
calculos = sorted(calculos, key = lambda x: [int(k) if k.isdigit() else k for k in re.split('([0-9]+)', x.stem)])

doc_calculos = Document(f'{os.path.dirname(__file__)}\MODELO_CALCULOS.docx')

for i, image in enumerate(calculos):
    picture = str(image)
    p = doc_calculos.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(picture, width=Inches(11.3), height=Inches(6.6))
    name_poste = os.path.split(image)
    name = (name_poste[1].split('.'))[0]
    doc_calculos.add_paragraph(
        f'POSTE {name}').paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc_calculos.save(f'{os.path.dirname(__file__)}\CALCULOS_TOTAL.docx')

convert(f'{os.path.dirname(__file__)}\CALCULOS_TOTAL.docx',
        f'{path_calcu}\{city}_CALCULOS.pdf')


def cal():
    return True
