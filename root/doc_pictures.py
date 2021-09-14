# -*- coding: utf-8 -*-
"""
Created on Wed Sep  8 11:35:01 2021

@author: heitor

RELATÓRIO PDF
"""

from docx import Document
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
#import tkinter as tk
from docx2pdf import convert
from pathlib import Path
from docx.shared import Inches
import root as rt
import document as dc

path_fotos = list(rt.getInput())[1]
city = list(rt.getInput())[2].upper()
path = Path(path_fotos)
fotos = list(path.glob('*.jpg'))

relatorio = str(
    f'{os.path.dirname(__file__)}\RELATORIO FOTOGRAFICO MODELO A4.docx')
doc_fotos = Document(relatorio)
dc.doc(doc_fotos)
document = Document(f'{os.path.dirname(__file__)}\FOTOS.docx')

for i, image in enumerate(fotos):
    picture = os.path.join(path_fotos, image)
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(picture, width=Inches(10), height=Inches(6))

document.save(f'{os.path.dirname(__file__)}\FOTOS_TOTAL.docx')

convert(f'{os.path.dirname(__file__)}\FOTOS_TOTAL.docx',
        f'{path_fotos}\{city}_FOTOS.pdf')


def pic():
    return True
